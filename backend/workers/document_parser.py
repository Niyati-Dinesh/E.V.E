"""
Document Parser Utility for E.V.E. Workers
Extracts text from various file formats: .docx, .pdf, .txt, code files
"""

import base64
import io
from docx import Document
from PyPDF2 import PdfReader

def parse_document(filename: str, base64_content: str) -> str:
    """
    Parse document and extract readable text
    
    Args:
        filename: Name of the file (determines parsing method)
        base64_content: Base64 encoded file content
    
    Returns:
        Extracted text content
    """
    
    # Decode base64
    try:
        file_bytes = base64.b64decode(base64_content)
    except Exception as e:
        return f"❌ Failed to decode file: {e}"
    
    file_lower = filename.lower()
    
    # Parse Word documents (.docx)
    if file_lower.endswith('.docx'):
        try:
            doc = Document(io.BytesIO(file_bytes))
            text_content = []
            
            # Extract all paragraphs
            for para in doc.paragraphs:
                if para.text.strip():
                    text_content.append(para.text)
            
            # Extract tables
            for table in doc.tables:
                for row in table.rows:
                    row_data = [cell.text for cell in row.cells]
                    text_content.append(" | ".join(row_data))
            
            return "\n".join(text_content) if text_content else "❌ Document appears to be empty"
            
        except Exception as e:
            return f"❌ Failed to parse Word document: {e}"
    
    # Parse PDF documents
    elif file_lower.endswith('.pdf'):
        try:
            pdf_file = io.BytesIO(file_bytes)
            pdf_reader = PdfReader(pdf_file)
            
            text_content = []
            for page_num, page in enumerate(pdf_reader.pages, 1):
                page_text = page.extract_text()
                if page_text.strip():
                    text_content.append(f"--- Page {page_num} ---\n{page_text}")
            
            return "\n\n".join(text_content) if text_content else "❌ PDF appears to be empty or unreadable"
            
        except Exception as e:
            return f"❌ Failed to parse PDF: {e}"
    
    # Parse text files (.txt, .md, .json, .csv, code files)
    elif any(file_lower.endswith(ext) for ext in ['.txt', '.md', '.json', '.csv', '.py', '.js', '.java', '.cpp', '.html', '.css', '.xml', '.yaml', '.yml']):
        try:
            text_content = file_bytes.decode('utf-8')
            return text_content
        except UnicodeDecodeError:
            try:
                # Try different encoding
                text_content = file_bytes.decode('latin-1')
                return text_content
            except Exception as e:
                return f"❌ Failed to decode text file: {e}"
    
    # Unsupported file type
    else:
        return f"⚠️ File type not supported for text extraction: {filename}\n(Supported: .docx, .pdf, .txt, .md, .json, .csv, code files)"

def process_uploaded_files(files: list) -> str:
    """
    Process a list of uploaded files and return formatted content
    
    Args:
        files: List of dicts with 'filename', 'content' (base64), 'mime_type'
    
    Returns:
        Formatted string with all file contents
    """
    if not files:
        return ""
    
    result = ["\n\n📎 **Uploaded Files:**\n"]
    
    for idx, file in enumerate(files, 1):
        filename = file.get('filename', f'file_{idx}')
        base64_content = file.get('content', '')
        size = file.get('size', 0)
        
        result.append(f"\n### File {idx}: {filename} ({format_size(size)})\n")
        
        # Parse the document
        parsed_content = parse_document(filename, base64_content)
        
        # Add content with proper formatting
        result.append("**Content:**\n```")
        result.append(parsed_content[:3000])  # Limit to first 3000 chars to avoid context overflow
        if len(parsed_content) > 3000:
            result.append("\n\n... (content truncated, file is large)")
        result.append("```\n")
    
    return "\n".join(result)

def format_size(size_bytes: int) -> str:
    """Format file size in human-readable format"""
    if size_bytes < 1024:
        return f"{size_bytes} bytes"
    elif size_bytes < 1024 * 1024:
        return f"{size_bytes / 1024:.1f} KB"
    else:
        return f"{size_bytes / (1024 * 1024):.1f} MB"

# Test function
if __name__ == "__main__":
    # Test with sample data
    import sys
    print("Document Parser Utility")
    print("=" * 50)
    print("Supported formats:")
    print("  - Word documents (.docx)")
    print("  - PDF files (.pdf)")
    print("  - Text files (.txt, .md, .json, .csv)")
    print("  - Code files (.py, .js, .java, .cpp, etc.)")
    print("\nReady to parse documents!")
