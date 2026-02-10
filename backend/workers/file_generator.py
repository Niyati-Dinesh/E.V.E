"""Generate PDF and DOCX files from markdown/text"""
from io import BytesIO
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
from reportlab.lib.enums import TA_LEFT, TA_CENTER
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont

try:
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    Document = None

try:
    import markdown
except ImportError:
    markdown = None


def markdown_to_text(md_text: str) -> str:
    """Convert markdown to plain text"""
    if markdown:
        import re
        # Simple markdown to text conversion
        text = md_text
        # Headers
        text = re.sub(r'^###\s+(.+)$', r'\1', text, flags=re.MULTILINE)
        text = re.sub(r'^##\s+(.+)$', r'\n=== \1 ===\n', text, flags=re.MULTILINE)
        text = re.sub(r'^#\s+(.+)$', r'\n=== \1 ===\n', text, flags=re.MULTILINE)
        # Bold/italic
        text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
        text = re.sub(r'\*(.+?)\*', r'\1', text)
        # Code blocks
        text = re.sub(r'```(\w*)\n(.+?)```', r'\n--- CODE ---\n\2\n--- END CODE ---\n', text, flags=re.DOTALL)
        # Inline code
        text = re.sub(r'`(.+?)`', r'\1', text)
        return text
    return md_text


def generate_pdf(content: str, title: str = "Document") -> tuple[BytesIO, str]:
    """Generate PDF from content, returns (file_data, filename)"""
    buffer = BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=letter,
                          rightMargin=72, leftMargin=72,
                          topMargin=72, bottomMargin=18)

    styles = getSampleStyleSheet()

    # Custom styles
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Heading1'],
        fontSize=24,
        spaceAfter=30,
        alignment=TA_CENTER,
        textColor='#2c3e50'
    )

    heading_style = ParagraphStyle(
        'CustomHeading',
        parent=styles['Heading2'],
        fontSize=16,
        spaceAfter=12,
        spaceBefore=20,
        textColor='#34495e'
    )

    normal_style = ParagraphStyle(
        'CustomNormal',
        parent=styles['Normal'],
        fontSize=11,
        spaceAfter=10,
        leading=14,
        textColor='#2c3e50'
    )

    # Build story
    story = []
    story.append(Paragraph(title, title_style))
    story.append(Spacer(1, 0.2 * inch))

    # Convert markdown-style content to paragraphs
    lines = content.split('\n')
    current_paragraph = ''

    for line in lines:
        line = line.strip()
        if not line:
            if current_paragraph:
                # Clean up markdown
                clean = current_paragraph.replace('**', '').replace('*', '').replace('`', '')
                story.append(Paragraph(clean, normal_style))
                story.append(Spacer(1, 6))
                current_paragraph = ''
            continue

        if line.startswith('#'):
            if current_paragraph:
                clean = current_paragraph.replace('**', '').replace('*', '').replace('`', '')
                story.append(Paragraph(clean, normal_style))
                story.append(Spacer(1, 6))
                current_paragraph = ''
            # Header
            header_text = line.lstrip('#').strip().replace('**', '').replace('*', '')
            story.append(Paragraph(header_text, heading_style))
            story.append(Spacer(1, 6))
        elif line.startswith('-') or line.startswith('*') or line.startswith('1.'):
            if current_paragraph:
                clean = current_paragraph.replace('**', '').replace('*', '').replace('`', '')
                story.append(Paragraph(clean, normal_style))
                story.append(Spacer(1, 6))
                current_paragraph = ''
            bullet = line.lstrip('-*0123456789.').strip()
            clean = bullet.replace('**', '').replace('*', '').replace('`', '')
            story.append(Paragraph(f'• {clean}', normal_style))
        else:
            current_paragraph += ' ' + line if current_paragraph else line

    if current_paragraph:
        clean = current_paragraph.replace('**', '').replace('*', '').replace('`', '')
        story.append(Paragraph(clean, normal_style))

    doc.build(story)
    buffer.seek(0)

    filename = f"{title.replace(' ', '_')[:30]}.pdf"
    return buffer, filename


def generate_docx(content: str, title: str = "Document") -> tuple[BytesIO, str]:
    """Generate DOCX from content, returns (file_data, filename)"""
    if not Document:
        raise ImportError("python-docx not installed. Run: pip install python-docx")

    buffer = BytesIO()
    doc = Document()

    # Title
    title_para = doc.add_paragraph(title.replace('#', '').strip())
    title_para.runs[0].bold = True
    title_para.runs[0].font.size = Pt(18)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Content
    lines = content.split('\n')
    for line in lines:
        line = line.strip()
        if not line:
            doc.add_paragraph()
            continue

        if line.startswith('###'):
            p = doc.add_paragraph(line.lstrip('#').strip().replace('**', ''))
            p.runs[0].font.size = Pt(14)
            p.runs[0].bold = True
        elif line.startswith('##'):
            p = doc.add_paragraph(line.lstrip('#').strip().replace('**', ''))
            p.runs[0].font.size = Pt(16)
            p.runs[0].bold = True
        elif line.startswith('#'):
            p = doc.add_paragraph(line.lstrip('#').strip().replace('**', ''))
            p.runs[0].font.size = Pt(18)
            p.runs[0].bold = True
        elif line.startswith('-') or line.startswith('*'):
            p = doc.add_paragraph(line.lstrip('-*').strip().replace('**', ''))
            p.style = 'List Bullet'
        elif line.startswith(('1.', '2.', '3.', '4.', '5.', '6.', '7.', '8.', '9.')):
            p = doc.add_paragraph(line[3:].strip().replace('**', ''))
            p.style = 'List Number'
        else:
            p = doc.add_paragraph(line.replace('**', '').replace('`', ''))

    doc.save(buffer)
    buffer.seek(0)

    filename = f"{title.replace(' ', '_')[:30]}.docx"
    return buffer, filename


def generate_file(content: str, file_type: str = "pdf", title: str = "Document") -> tuple[BytesIO, str, str]:
    """
    Generate file from content
    Returns: (file_data, filename, mime_type)
    """
    file_type = file_type.lower()

    if file_type == "pdf":
        data, filename = generate_pdf(content, title)
        return data, filename, "application/pdf"
    elif file_type in ["docx", "doc"]:
        data, filename = generate_docx(content, title)
        return data, filename, "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    else:
        raise ValueError(f"Unsupported file type: {file_type}")